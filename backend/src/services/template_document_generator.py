"""
Template-based document generator for ATS-friendly PDF and DOCX resumes.
Uses HTML templates with Jinja2 for consistent, professional formatting.

This replaces the manual ReportLab-based generation with a template-driven approach:
- HTML templates → WeasyPrint → PDF (ATS-optimized)
- Template data → python-docx → DOCX (editable)
- 100% format consistency between PDF and DOCX
"""

from __future__ import annotations

import io
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader, select_autoescape
from loguru import logger
from weasyprint import HTML, CSS

try:
    from ..models.resume import Resume
except ImportError:
    from models.resume import Resume


class TemplateDocumentGenerator:
    """
    Template-based document generator for PDF and DOCX resumes.

    Features:
    - Jinja2 HTML templates for consistent styling
    - WeasyPrint for HTML → PDF conversion (better ATS compatibility than xhtml2pdf)
    - python-docx for editable Word documents
    - Automatic skill categorization
    - Clean, single-column ATS-friendly layout
    """

    def __init__(self, template_dir: Optional[str] = None):
        """
        Initialize the template-based document generator.

        Args:
            template_dir: Path to templates directory (defaults to backend/templates)
        """
        if template_dir is None:
            # Default to backend/templates directory
            current_file = Path(__file__).resolve()
            backend_root = current_file.parent.parent.parent  # Go up to backend/
            template_dir = str(backend_root / "templates")

        self.template_dir = template_dir

        # Initialize Jinja2 environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(template_dir),
            autoescape=select_autoescape(['html', 'xml']),
            trim_blocks=True,
            lstrip_blocks=True
        )

        logger.info(f"Initialized TemplateDocumentGenerator with templates from: {template_dir}")

    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """
        Categorize skills into logical groups for better organization.

        Args:
            skills: List of skill strings

        Returns:
            Dictionary mapping category names to skill lists
        """
        # Skill categorization keywords
        categories = {
            "Languages": ["python", "javascript", "typescript", "java", "c++", "c#", "go", "rust", "ruby", "php", "kotlin", "swift", "scala"],
            "Frontend": ["react", "vue", "angular", "svelte", "html", "css", "sass", "tailwind", "bootstrap", "webpack", "vite"],
            "Backend": ["node.js", "express", "fastapi", "django", "flask", "spring", "asp.net", ".net", "rails"],
            "Databases": ["postgresql", "mysql", "mongodb", "redis", "elasticsearch", "sql", "nosql", "dynamodb", "cassandra"],
            "Cloud & DevOps": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins", "github actions", "ci/cd", "linux"],
            "Tools & Other": []
        }

        # Categorize each skill
        categorized = {cat: [] for cat in categories.keys()}
        uncategorized = []

        for skill in skills:
            skill_lower = skill.lower()
            matched = False

            for category, keywords in categories.items():
                if category == "Tools & Other":
                    continue

                for keyword in keywords:
                    if keyword in skill_lower or skill_lower in keyword:
                        categorized[category].append(skill)
                        matched = True
                        break

                if matched:
                    break

            if not matched:
                uncategorized.append(skill)

        # Add uncategorized to "Tools & Other"
        categorized["Tools & Other"] = uncategorized

        # Remove empty categories
        categorized = {k: v for k, v in categorized.items() if v}

        return categorized

    def _prepare_template_data(self, resume: Resume) -> Dict[str, Any]:
        """
        Prepare resume data for template rendering.

        Args:
            resume: Resume model instance

        Returns:
            Dictionary with all template variables
        """
        # Convert resume to dict
        data = resume.model_dump()

        # Categorize skills if present
        if resume.skills and len(resume.skills) > 0:
            data['skillsGrouped'] = self._categorize_skills(resume.skills)

        # Format dates
        for exp in data.get('experience', []):
            if exp.get('endDate') and exp['endDate'].lower() in ['present', 'current']:
                exp['endDate'] = 'Present'

        # Add metadata
        data['generatedDate'] = 'Generated with Resume Tailor AI'

        return data

    # ==================== PDF GENERATION (WeasyPrint) ====================

    def generate_pdf(self, resume: Resume) -> bytes:
        """
        Generate ATS-friendly PDF from HTML template using WeasyPrint.

        WeasyPrint advantages over xhtml2pdf:
        - Better CSS support (CSS3, flexbox, etc.)
        - More accurate rendering
        - Better font handling
        - Cleaner output for ATS systems

        Args:
            resume: Resume model instance

        Returns:
            PDF file as bytes

        Raises:
            Exception: If PDF generation fails
        """
        try:
            logger.info(f"Generating PDF from template for: {resume.personalInfo.name}")

            # Prepare template data
            template_data = self._prepare_template_data(resume)

            # Load and render HTML template
            template = self.jinja_env.get_template('resume_template.html')
            html_content = template.render(**template_data)

            # Generate PDF using WeasyPrint
            pdf_bytes = HTML(string=html_content, base_url=self.template_dir).write_pdf()

            logger.info(f"PDF generated successfully: {len(pdf_bytes)} bytes")
            return pdf_bytes

        except Exception as e:
            logger.error(f"Failed to generate PDF: {str(e)}")
            raise Exception(f"PDF generation failed: {str(e)}")

    # ==================== DOCX GENERATION (python-docx) ====================

    def generate_docx(self, resume: Resume) -> bytes:
        """
        Generate editable DOCX resume with formatting matching the PDF template.

        Args:
            resume: Resume model instance

        Returns:
            DOCX file as bytes

        Raises:
            Exception: If DOCX generation fails
        """
        try:
            logger.info(f"Generating DOCX for: {resume.personalInfo.name}")

            # Create document
            doc = Document()

            # Set document margins (0.75 inches to match HTML template)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # ==== HEADER - Name and Contact Info ====
            self._add_docx_header(doc, resume.personalInfo)

            # ==== PROFESSIONAL SUMMARY ====
            if resume.personalInfo.summary:
                self._add_docx_section(doc, "PROFESSIONAL SUMMARY", resume.personalInfo.summary)

            # ==== PROFESSIONAL EXPERIENCE ====
            if resume.experience and len(resume.experience) > 0:
                self._add_docx_experience(doc, resume.experience)

            # ==== EDUCATION ====
            if resume.education and len(resume.education) > 0:
                self._add_docx_education(doc, resume.education)

            # ==== TECHNICAL SKILLS ====
            if resume.skills and len(resume.skills) > 0:
                self._add_docx_skills(doc, resume.skills)

            # ==== PROJECTS ====
            if resume.projects and len(resume.projects) > 0:
                self._add_docx_projects(doc, resume.projects)

            # ==== CERTIFICATIONS ====
            if resume.certifications and len(resume.certifications) > 0:
                self._add_docx_certifications(doc, resume.certifications)

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_bytes = buffer.read()

            logger.info(f"DOCX generated successfully: {len(docx_bytes)} bytes")
            return docx_bytes

        except Exception as e:
            logger.error(f"Failed to generate DOCX: {str(e)}")
            raise Exception(f"DOCX generation failed: {str(e)}")

    # ==================== DOCX HELPER METHODS ====================

    def _add_docx_header(self, doc: Document, personal_info):
        """Add header with name and contact information"""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal_info.name)
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(26, 26, 26)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_parts = []
        if personal_info.email:
            contact_parts.append(personal_info.email)
        if personal_info.phone:
            contact_parts.append(personal_info.phone)
        if personal_info.location:
            contact_parts.append(personal_info.location)
        if personal_info.linkedin:
            contact_parts.append("LinkedIn")
        if personal_info.github:
            contact_parts.append("GitHub")
        if personal_info.website:
            contact_parts.append("Portfolio")

        if contact_parts:
            contact_para = doc.add_paragraph(" • ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.runs[0]
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = RGBColor(75, 85, 99)

        # Add horizontal line
        doc.add_paragraph()

    def _add_docx_section(self, doc: Document, title: str, content: Optional[str] = None):
        """Add a section with title and optional content"""
        # Section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(37, 99, 235)  # #2563eb

        # Content paragraph if provided
        if content:
            content_para = doc.add_paragraph(content)
            content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            content_para.paragraph_format.space_after = Pt(12)

        doc.add_paragraph()  # Spacing

    def _add_docx_experience(self, doc: Document, experience: List):
        """Add professional experience section"""
        self._add_docx_section(doc, "PROFESSIONAL EXPERIENCE")

        for exp in experience:
            # Position title
            position_para = doc.add_paragraph()
            position_run = position_para.add_run(exp.position)
            position_run.font.size = Pt(12)
            position_run.font.bold = True

            # Company name
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(exp.company)
            company_run.font.size = Pt(11)
            company_run.font.bold = True
            company_run.font.color.rgb = RGBColor(75, 85, 99)

            # Date range and location
            meta_parts = []
            if exp.startDate:
                date_str = f"{exp.startDate}"
                if exp.endDate:
                    date_str += f" - {exp.endDate if exp.endDate else 'Present'}"
                meta_parts.append(date_str)
            if exp.location:
                meta_parts.append(exp.location)

            if meta_parts:
                meta_para = doc.add_paragraph(" • ".join(meta_parts))
                meta_run = meta_para.runs[0]
                meta_run.font.size = Pt(10)
                meta_run.font.color.rgb = RGBColor(107, 114, 128)
                meta_run.italic = True

            # Bullets
            if exp.bullets:
                for bullet in exp.bullets:
                    bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                    bullet_para.paragraph_format.left_indent = Inches(0.25)

            doc.add_paragraph()  # Spacing between experiences

    def _add_docx_education(self, doc: Document, education: List):
        """Add education section"""
        self._add_docx_section(doc, "EDUCATION")

        for edu in education:
            # Degree
            degree_str = edu.degree
            if edu.field:
                degree_str += f" in {edu.field}"
            if edu.gpa:
                degree_str += f" - GPA: {edu.gpa}"

            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(degree_str)
            degree_run.font.size = Pt(11.5)
            degree_run.font.bold = True

            # Institution
            inst_para = doc.add_paragraph(edu.institution)
            inst_run = inst_para.runs[0]
            inst_run.font.color.rgb = RGBColor(75, 85, 99)

            # Dates
            if edu.startDate:
                date_str = edu.startDate
                if edu.endDate:
                    date_str += f" - {edu.endDate}"
                date_para = doc.add_paragraph(date_str)
                date_run = date_para.runs[0]
                date_run.font.size = Pt(10)
                date_run.font.color.rgb = RGBColor(107, 114, 128)
                date_run.italic = True

            # Achievements
            if edu.achievements:
                for achievement in edu.achievements:
                    ach_para = doc.add_paragraph(achievement, style='List Bullet 2')
                    ach_para.paragraph_format.left_indent = Inches(0.25)

            doc.add_paragraph()  # Spacing

    def _add_docx_skills(self, doc: Document, skills: List[str]):
        """Add technical skills section"""
        self._add_docx_section(doc, "TECHNICAL SKILLS")

        # Categorize skills
        categorized = self._categorize_skills(skills)

        if categorized:
            for category, skill_list in categorized.items():
                skill_para = doc.add_paragraph()

                # Category name (bold)
                cat_run = skill_para.add_run(f"{category}: ")
                cat_run.font.bold = True
                cat_run.font.color.rgb = RGBColor(55, 65, 81)

                # Skills list
                skills_run = skill_para.add_run(", ".join(skill_list))
                skills_run.font.color.rgb = RGBColor(75, 85, 99)
        else:
            # Simple list if no categorization
            skills_para = doc.add_paragraph(" • ".join(skills))

        doc.add_paragraph()  # Spacing

    def _add_docx_projects(self, doc: Document, projects: List):
        """Add projects section"""
        self._add_docx_section(doc, "PROJECTS")

        for project in projects:
            # Project name
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(project.name)
            name_run.font.size = Pt(11.5)
            name_run.font.bold = True

            if project.url:
                url_run = name_para.add_run(f" [{project.url}]")
                url_run.font.size = Pt(10)
                url_run.font.color.rgb = RGBColor(29, 78, 216)

            # Description
            if project.description:
                desc_para = doc.add_paragraph(project.description)
                desc_para.paragraph_format.space_after = Pt(4)

            # Technologies
            if project.technologies:
                tech_para = doc.add_paragraph(f"Technologies: {', '.join(project.technologies)}")
                tech_run = tech_para.runs[0]
                tech_run.font.size = Pt(10)
                tech_run.font.color.rgb = RGBColor(107, 114, 128)
                tech_run.italic = True

            doc.add_paragraph()  # Spacing

    def _add_docx_certifications(self, doc: Document, certifications: List):
        """Add certifications section"""
        self._add_docx_section(doc, "CERTIFICATIONS")

        for cert in certifications:
            cert_para = doc.add_paragraph()

            # Certification name (bold)
            name_run = cert_para.add_run(cert.name)
            name_run.font.bold = True

            # Issuer
            if cert.issuer:
                issuer_run = cert_para.add_run(f" - {cert.issuer}")
                issuer_run.font.color.rgb = RGBColor(75, 85, 99)

            # Date
            if cert.date:
                date_run = cert_para.add_run(f" ({cert.date})")
                date_run.font.size = Pt(10)
                date_run.font.color.rgb = RGBColor(107, 114, 128)

        doc.add_paragraph()  # Spacing


# ==================== FACTORY FUNCTION ====================

def get_template_generator(template_dir: Optional[str] = None) -> TemplateDocumentGenerator:
    """
    Factory function to get a TemplateDocumentGenerator instance.

    Args:
        template_dir: Optional path to templates directory

    Returns:
        TemplateDocumentGenerator instance
    """
    return TemplateDocumentGenerator(template_dir=template_dir)
